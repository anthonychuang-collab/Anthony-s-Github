# -*- coding: utf-8 -*-
"""F 班系統測試套件（零依賴，python3 tests/test_fban.py）。"""
import sys, os, traceback
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fban.config import Config, Person, HeadStaff
from fban.codes import CodeBook, leading_floor
from fban.convert import assign_rest, _place_li, convert_person, convert_foreign_code
from fban.coverage import check_coverage, check_labor
from fban.fillin import auto_fill
from fban import writer, read_fban
import tempfile, os as _os

_PASS = _FAIL = 0
def check(name, cond, detail=""):
    global _PASS, _FAIL
    if cond:
        _PASS += 1; print(f"  ✅ {name}")
    else:
        _FAIL += 1; print(f"  ❌ {name}  {detail}")
def eq(name, got, want):
    check(name, got == want, f"got={got!r} want={want!r}")

CODE_MAP = [
    {"T班原始碼": "Di", "F班顯示碼": "D4x", "班別大類": "D白"},
    {"T班原始碼": "2Di", "F班顯示碼": "D4x", "班別大類": "D白"},
    {"T班原始碼": "3Di", "F班顯示碼": "D4x", "班別大類": "D白"},
    {"T班原始碼": "D", "F班顯示碼": "D4x", "班別大類": "D白"},
    {"T班原始碼": "D6", "F班顯示碼": "D6x", "班別大類": "D白"},
    {"T班原始碼": "E", "F班顯示碼": "Ex", "班別大類": "E小夜"},
    {"T班原始碼": "N", "F班顯示碼": "Nx", "班別大類": "N大夜"},
]
COLOR_RULES = [
    {"區塊": "護理", "班別大類": "D白", "樓層": "2F", "RGB(ARGB)": "FF70AD47"},
    {"區塊": "護理", "班別大類": "D白", "樓層": "3F", "RGB(ARGB)": "FF5B9BD5"},
    {"區塊": "護理", "班別大類": "D白", "樓層": "5F", "RGB(ARGB)": "FFFF2F92"},
    {"區塊": "護理", "班別大類": "E小夜", "樓層": "*", "RGB(ARGB)": "FFED7D31"},
    {"區塊": "護理", "班別大類": "N大夜", "樓層": "*", "RGB(ARGB)": "FF44546A"},
    {"區塊": "台籍照服", "班別大類": "D白", "樓層": "5F", "RGB(ARGB)": "FFFF2F92"},
    {"區塊": "外籍照服", "班別大類": "D白", "樓層": "3F", "RGB(ARGB)": "FF5B9BD5"},
    {"區塊": "外籍照服", "班別大類": "N大夜", "樓層": "5F", "RGB(ARGB)": "FFD9D9D9"},
]
def make_cfg(**settings):
    cfg = Config(); cfg.code_map = CODE_MAP; cfg.color_rules = COLOR_RULES
    cfg.settings = {"每14天最少例假": 2, "西元年": 2026, "月份": 8, "每班最低人力": 7}
    cfg.settings.update(settings)
    return cfg
def days_from(codes):
    return {"name": "X", "account": "1", "days": {i+1: c for i, c in enumerate(codes)}}


def test_codebook_nursing():
    cb = CodeBook(CODE_MAP)
    eq("Di→D4x/D白", cb.lookup("Di")[:2], ("D4x", "D白"))
    eq("2Di 樓層=2F", cb.lookup("2Di")[2], "2F")
    eq("3Di 樓層=3F", cb.lookup("3Di")[2], "3F")
    eq("E→Ex", cb.lookup("E")[:2], ("Ex", "E小夜"))
    eq("N→Nx", cb.lookup("N")[:2], ("Nx", "N大夜"))

def test_codebook_foreign_both_orders():
    cb = CodeBook(CODE_MAP)
    eq("D3a→Dx", cb.lookup("D3a")[0], "Dx")
    eq("D3a 樓層=3F", cb.lookup("D3a")[2], "3F")
    eq("3Da→Dx", cb.lookup("3Da")[0], "Dx")
    eq("N5a→Nx", cb.lookup("N5a")[0], "Nx")
    eq("N5a 樓層=5F", cb.lookup("N5a")[2], "5F")
    eq("D5b→Dx", cb.lookup("D5b")[0], "Dx")

def test_codebook_pt_and_kitchen():
    cb = CodeBook(CODE_MAP)
    eq("Ep→其他", cb.lookup("Ep")[1], "其他")
    eq("Dp→其他", cb.lookup("Dp")[1], "其他")
    eq("C→其他", cb.lookup("C")[1], "其他")
    eq("K→其他", cb.lookup("K")[1], "其他")

def test_leading_floor():
    eq("2Di 前綴=2F", leading_floor("2Di"), "2F")
    eq("Di 無前綴", leading_floor("Di"), None)

def test_assign_rest_quota_counts():
    rest = [1, 4, 8, 12, 16, 20, 24, 28, 31]
    labels = assign_rest(rest, 2, None, {"例": 5, "休": 4, "國": 0}, set())
    eq("配額5例→恰5例", sum(1 for v in labels.values() if v == "例"), 5)
    eq("其餘為休", sum(1 for v in labels.values() if v == "休"), 4)

def test_assign_rest_quota_guo_on_holiday():
    labels = assign_rest([1, 5, 10, 15, 20, 25], 2, None, {"例": 4, "休": 0, "國": 2}, {5, 20})
    check("國定假日確切日→標國", labels.get(5) == "國" and labels.get(20) == "國", labels)
    eq("國恰2天", sum(1 for v in labels.values() if v == "國"), 2)

def test_assign_rest_no_quota_alternate():
    labels = assign_rest([2, 5, 9, 13], 2, None, None, None)
    check("無配額≥2例", sum(1 for v in labels.values() if v == "例") >= 2, labels)

def test_place_li_spread():
    li = _place_li([1, 3, 6, 9, 12, 15, 18, 21, 24, 27, 30], 5, 2)
    eq("place_li 回傳5個", len(li), 5)
    check("例假有分散(跨月首尾)", min(li) <= 6 and max(li) >= 24, sorted(li))

def test_convert_nursing_codes_and_color():
    cfg = make_cfg()
    conv = convert_person(Person(name="甲", block="護理"),
                          days_from(["Di", "2Di", "3Di", "E", "N", "R", "R"]), CodeBook(CODE_MAP), cfg)
    d = conv["days"]
    eq("Di→D4x", d[1]["code"], "D4x")
    eq("Di 5F紅", d[1]["color"], "FFFF2F92")
    eq("2Di 2F綠", d[2]["color"], "FF70AD47")
    eq("3Di 3F藍", d[3]["color"], "FF5B9BD5")
    eq("E→Ex橘", (d[4]["code"], d[4]["color"]), ("Ex", "FFED7D31"))
    eq("N→Nx灰", (d[5]["code"], d[5]["color"]), ("Nx", "FF44546A"))
    check("R→例或休", d[6]["code"] in ("例", "休", "國"), d[6])

def test_white_code_by_block_and_allowed():
    cfg = make_cfg(); cb = CodeBook(CODE_MAP)
    eq("護理白→D4x", convert_person(Person(name="護", block="護理"), days_from(["D"]), cb, cfg)["days"][1]["code"], "D4x")
    eq("台籍白→D5x", convert_person(Person(name="台", block="台籍照服"), days_from(["D"]), cb, cfg)["days"][1]["code"], "D5x")
    eq("台籍勾D6x→D6x", convert_person(Person(name="台6", block="台籍照服", allowed={"D6x","Ex","Nx"}), days_from(["D"]), cb, cfg)["days"][1]["code"], "D6x")

def test_person_head_name():
    eq("人頭核章=牌照持有人", Person(name="陳詡善", block="護理", stamp_name="陳淑萍").record_name, "陳淑萍")
    eq("本人核章=自己", Person(name="甲", block="護理").record_name, "甲")

def test_coverage_flags_shortfall():
    cfg = make_cfg()
    p = convert_person(Person(name="甲", block="護理"),
                       {"name":"甲","account":"1","days":{d:"Di" for d in range(1,4)}}, CodeBook(CODE_MAP), cfg)
    issues, _ = check_coverage([p], 3, cfg)
    check("偵測白班<7", any("僅" in s for s in issues), issues[:3])
    check("偵測缺護理", any("護理" in s for s in issues), issues[:3])

def test_labor_pt_filter():
    cfg = make_cfg(); cb = CodeBook(CODE_MAP)
    pt = convert_person(Person(name="PT員", block="護理"), days_from(["E","E","E"]+["/"]*28), cb, cfg)
    check("PT不報例假", not any("PT員" in s for s in check_labor([pt], cfg)), "")

def test_labor_detects_reverse_order():
    cfg = make_cfg(); cb = CodeBook(CODE_MAP)
    p = convert_person(Person(name="逆", block="護理"), days_from(["N","Di"]+["Di"]*10), cb, cfg)
    check("偵測逆排 大夜→白", any("順排" in s for s in check_labor([p], cfg)), "")

def test_fill_availability_gao():
    cfg = make_cfg()
    cfg.head_pool = [HeadStaff(name="高偲芸", category="護理", avail_from=(115,8,20))]
    heads, _ = auto_fill([], 31, cfg)
    gao = next((h for h in heads if h["name"] == "高偲芸"), None)
    if gao:
        eq("高偲芸8/20前不上班", [d for d in range(1,20) if gao["days"][d]["is_work"]], [])

def test_fill_meets_minimums():
    cfg = make_cfg(白班目標人數=1, 小夜目標人數=1, 大夜目標人數=1)
    cfg.head_pool = ([HeadStaff(name=f"護{i}", category="護理") for i in range(6)] +
                     [HeadStaff(name=f"台{i}", category="台籍照服") for i in range(6)])
    heads, _ = auto_fill([], 5, cfg)
    _, counts = check_coverage(heads, 5, cfg)
    ok = all(counts[d][s].get("護理",0) >= 1 for d in range(1,6) for s in ["D白","E小夜","N大夜"])
    check("補人頭後每班護理≥1", ok, "")

def test_fill_head_no_labor_violation():
    cfg = make_cfg(白班目標人數=1, 小夜目標人數=1, 大夜目標人數=1)
    cfg.head_pool = [HeadStaff(name=f"護{i}", category="護理") for i in range(8)] + \
                    [HeadStaff(name=f"台{i}", category="台籍照服") for i in range(8)]
    heads, _ = auto_fill([], 31, cfg)
    eq("人頭班0勞基法違規", check_labor(heads, cfg), [])

def test_convert_foreign_code_helper():
    eq("2Da→Dax(舊helper)", convert_foreign_code("2Da"), "Dax")
    eq("3Nb→Nbx(舊helper)", convert_foreign_code("3Nb"), "Nbx")

def _mk(name, block, record, per_day):
    days = {}
    for d in range(1, 4):
        cat, fl = per_day.get(d, (None, None))
        days[d] = {"code": cat, "cat": cat, "floor": fl, "is_work": cat in ("D白","E小夜","N大夜")}
    return {"name": name, "record_name": record, "block": block, "shift_kind": "", "n_days": 3, "days": days}

def test_docgen_restraint():
    from fban import docgen
    conv = [_mk("護A","護理","牌照A",{1:("D白","2F"),2:("D白","3F")}),
            _mk("護B","護理","護B",{1:("N大夜",None)}),
            _mk("護C","護理","護C",{1:("E小夜",None)})]
    data = docgen.restraint_floor_data(conv, 3, prev_night="前月大夜")
    eq("2F白=牌照A", data[1]["2F"]["白班"], "牌照A")
    eq("第1天大夜=上月最後一天", data[1]["2F"]["大夜"], "前月大夜")
    eq("第2天大夜=班表第1天大夜(往前推一天)", data[2]["2F"]["大夜"], "護B")
    eq("小夜共用護C(當日)", data[1]["3F"]["小夜"], "護C")

def test_docgen_namecopy():
    from fban import docgen
    conv = [_mk("護A","護理","護A",{1:("D白","2F")}),
            _mk("台A","台籍照服","台A",{1:("D白","2F")}),
            _mk("台N","台籍照服","台N",{1:("E小夜","2F")}),
            _mk("外A","外籍照服","外A",{1:("N大夜","2F")})]
    a = docgen.namecopy_assignments(conv, 3)
    eq("2F白班護理", a["2F"]["1"]["白班護理"], "護A")
    eq("2F白班照服=台籍", a["2F"]["1"]["白班照服"], "台A")
    eq("2F夜班照服=台籍(非外籍)", a["2F"]["1"]["夜班照服"], "台N")


def _mk_head(name, block, record, per_day, is_head=False):
    p = _mk(name, block, record, per_day)
    p["is_head"] = is_head
    return p

def test_docgen_avoid_head_nurse():
    """同日同樓層兩位白班：責任護士取實際護理，跳過護理長(is_head)。"""
    from fban import docgen
    # 護理長排在前面，仍應取後面的實際責任護士
    conv = [_mk_head("護理長", "護理", "護理長", {1: ("D白", "2F")}, is_head=True),
            _mk_head("曾素靖", "護理", "曾素靖", {1: ("D白", "2F")}, is_head=False)]
    data = docgen.restraint_floor_data(conv, 3)
    eq("2F白避開護理長→曾素靖", data[1]["2F"]["白班"], "曾素靖")
    # 只有護理長在該樓白班時，仍要顯示護理長（不留白）
    conv2 = [_mk_head("護理長", "護理", "護理長", {1: ("D白", "3F")}, is_head=True)]
    data2 = docgen.restraint_floor_data(conv2, 3)
    eq("3F只剩護理長→仍取護理長", data2[1]["3F"]["白班"], "護理長")

def test_readfban_real_format():
    """機構原生版面：多分頁依月份挑、欄位靠標題對位、theme/RGB 底色判樓層、護理長班種D0。"""
    import openpyxl
    from openpyxl.styles import PatternFill
    cfg = make_cfg()
    wb = openpyxl.Workbook()
    wb.active.title = "115.08"          # 舊月份（不應被選到）
    wb.active["A1"] = "舊月份不選"
    ws = wb.create_sheet("115.09")      # 目標月份
    # 日期列：從第 11 欄起 1..30
    for d in range(1, 31):
        ws.cell(4, 10 + d, d)
    # 區塊表頭列
    hdr = {4: "帳號", 5: "核章人員", 6: "護理人員", 7: "班種"}
    for c, v in hdr.items():
        ws.cell(5, c, v)
    green = PatternFill("solid", fgColor="FF70AD47")   # 2F
    # 護理長(班種 D0)與實際護理，第1日都 2F 白
    ws.cell(6, 4, "R001"); ws.cell(6, 5, "顏欣盈"); ws.cell(6, 6, "顏欣盈"); ws.cell(6, 7, "D0")
    ws.cell(6, 11, "D4x").fill = green
    ws.cell(7, 4, "R190"); ws.cell(7, 6, "曾素靖")   # 核章空→用姓名
    ws.cell(7, 11, "D4x").fill = green
    tmp = _os.path.join(tempfile.gettempdir(), "test_real_F.xlsx")
    wb.save(tmp)
    conv, nd = read_fban.load(tmp, cfg, "115.09")
    eq("依月份挑到 115.09(30天)", nd, 30)
    names = {p["name"] for p in conv if p["block"] == "護理"}
    check("讀到護理兩人", names == {"顏欣盈", "曾素靖"}, f"got={names}")
    byname = {p["name"]: p for p in conv}
    eq("底色判樓層=2F", byname["曾素靖"]["days"][1]["floor"], "2F")
    eq("班種D0→is_head", byname["顏欣盈"]["is_head"], True)
    eq("核章空→回姓名", byname["曾素靖"]["record_name"], "曾素靖")
    from fban import docgen
    data = docgen.restraint_floor_data(conv, nd)
    eq("責任護士避開護理長→曾素靖", data[1]["2F"]["白班"], "曾素靖")


def _mk_full(name, block, rec, per, cfg):
    """建立含 color 的 converted person（供 writer 寫出）。"""
    days = {}
    for d in range(1, 32):
        c = per.get(d)
        cat = (c if c in ("例", "休", "國") else
               "D白" if c and c[0] == "D" else
               "E小夜" if c == "Ex" else "N大夜" if c == "Nx" else "空")
        info = {"code": c, "cat": cat, "floor": per.get(("fl", d)),
                "color": None, "is_work": c in ("D4x", "D5x", "Dx", "Ex", "Nx")}
        if info["is_work"]:
            if not info["floor"]:
                info["floor"] = "5F"
            info["color"] = cfg.color_for(block, cat, info["floor"])
        days[d] = info
    return {"name": name, "record_name": rec, "account": "X", "block": block,
            "shift_kind": "", "n_days": 31, "days": days}

def test_readfban_roundtrip():
    """US-7：F班寫出→讀回，人數/區塊/樓層/班別/核章 還原正確。"""
    cfg = make_cfg(週起始星期=6)
    conv = [
        _mk_full("顏欣盈", "護理", "顏欣盈",
                 {**{d: "D4x" for d in range(1, 32)}, **{("fl", d): "2F" for d in range(1, 32)}}, cfg),
        _mk_full("何承祐", "護理", "何承祐", {d: "Nx" for d in range(1, 32)}, cfg),
        _mk_full("陳詡善", "護理", "陳淑萍",  # 人頭：核章≠姓名
                 {**{d: "D4x" for d in range(1, 32)}, **{("fl", d): "3F" for d in range(1, 32)}}, cfg),
    ]
    tmp = _os.path.join(tempfile.gettempdir(), "test_rt_F.xlsx")
    writer.write(conv, 31, cfg, tmp, "115.08")
    conv2, nd = read_fban.load(tmp, cfg)
    eq("讀回天數=31", nd, 31)
    eq("讀回人數=3", len([p for p in conv2 if p["block"] == "護理"]), 3)
    byname = {p["name"]: p for p in conv2}
    eq("顏欣盈 第1日=D白", byname["顏欣盈"]["days"][1]["cat"], "D白")
    eq("顏欣盈 第1日樓層=2F(由底色)", byname["顏欣盈"]["days"][1]["floor"], "2F")
    eq("何承祐 第1日=N大夜", byname["何承祐"]["days"][1]["cat"], "N大夜")
    eq("陳詡善文件用名=陳淑萍(牌照持有人)", byname["陳詡善"]["record_name"], "陳淑萍")
    eq("陳詡善名冊欄=陳詡善(實際同仁)", byname["陳詡善"]["name"], "陳詡善")
    eq("陳詡善核章欄另存=陳淑萍", byname["陳詡善"]["stamp"], "陳淑萍")

def test_readfban_feeds_docgen():
    """US-7→US-8：讀回的資料能正確產生約束表指派。"""
    from fban import docgen
    cfg = make_cfg(週起始星期=6)
    conv = [
        _mk_full("顏欣盈", "護理", "顏欣盈",
                 {**{d: "D4x" for d in range(1, 32)}, **{("fl", d): "2F" for d in range(1, 32)}}, cfg),
        _mk_full("何承祐", "護理", "何承祐", {d: "Nx" for d in range(1, 32)}, cfg),
        _mk_full("黃安宇", "護理", "黃安宇", {d: "Ex" for d in range(1, 32)}, cfg),
    ]
    tmp = _os.path.join(tempfile.gettempdir(), "test_rt_F2.xlsx")
    writer.write(conv, 31, cfg, tmp, "115.08")
    conv2, nd = read_fban.load(tmp, cfg)
    data = docgen.restraint_floor_data(conv2, nd, prev_night="上月大夜")
    eq("2F白=顏欣盈", data[1]["2F"]["白班"], "顏欣盈")
    eq("第1天大夜=上月最後一天", data[1]["2F"]["大夜"], "上月大夜")
    eq("第2天大夜=何承祐(往前推一天)", data[2]["2F"]["大夜"], "何承祐")
    eq("小夜=黃安宇", data[1]["2F"]["小夜"], "黃安宇")


def test_namecopy_blank_31_in_30day_month():
    """9 月（30天）：照護表第 31 欄（及右側多餘欄）不得有姓名。"""
    try:
        import docx
    except Exception:
        print("  ⏭ 略過（無 python-docx）"); return
    from docx import Document
    import sys as _sys
    _sys.path.insert(0, _os.path.join(
        _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))), "docskills", "namecopy"))
    import importlib, json
    fcr = importlib.import_module("fill_care_record")
    # 造一個含 1..31 日欄 + 責任護士列的最小範本，並在第 31 欄預填殘留姓名
    d = Document()
    d.add_paragraph("115 年 9 月  姓名:")
    t = d.add_table(rows=2, cols=1 + 31)
    t.rows[0].cells[0].text = "日期"
    for day in range(1, 32):
        t.rows[0].cells[day].text = str(day)
    t.rows[1].cells[0].text = "責任護士 簽名"
    t.rows[1].cells[31].text = "殘留人"          # 第31欄殘留
    tmp = _os.path.join(tempfile.gettempdir(), "tpl_care.docx"); d.save(tmp)
    aj = _os.path.join(tempfile.gettempdir(), "asg.json")
    with open(aj, "w", encoding="utf-8") as f:
        json.dump({"2F": {str(x): {"白班護理": f"護{x}"} for x in range(1, 31)}}, f, ensure_ascii=False)
    out = _os.path.join(tempfile.gettempdir(), "care_out.docx")
    fcr.fill(tmp, aj, "2F", out)
    d2 = Document(out); row = d2.tables[0].rows[1]
    eq("第30欄有名", row.cells[30].text.strip(), "護30")
    eq("第31欄留白(9月無31)", row.cells[31].text.strip(), "")


def test_readfban_roundtrip_blocks():
    """迴歸：F班寫出→讀回，三個區塊要各自還原，不能全部落到護理。
    （writer 在表頭上一列寫的是「護理人員／台籍照服員／外籍照服員」全稱，
      read_fban 若只認短稱，會把每個人都當成護理，照護表的照服欄就會全空。）"""
    from fban import docgen
    cfg = make_cfg(週起始星期=6)
    conv = [
        _mk_full("護理甲", "護理", "護理甲",
                 {**{d: "D4x" for d in range(1, 32)}, **{("fl", d): "2F" for d in range(1, 32)}}, cfg),
        # 樓層須挑測試設定中有顏色規則者（台籍照服白=5F、外籍照服白=3F）
        _mk_full("台照乙", "台籍照服", "台照乙",
                 {**{d: "D5x" for d in range(1, 32)}, **{("fl", d): "5F" for d in range(1, 32)}}, cfg),
        _mk_full("外照丙", "外籍照服", "外照丙",
                 {**{d: "Dx" for d in range(1, 32)}, **{("fl", d): "3F" for d in range(1, 32)}}, cfg),
    ]
    tmp = _os.path.join(tempfile.gettempdir(), "test_rt_blocks.xlsx")
    writer.write(conv, 31, cfg, tmp, "115.08")
    conv2, nd = read_fban.load(tmp, cfg)
    got = {p["name"]: p["block"] for p in conv2}
    eq("護理甲 區塊=護理", got.get("護理甲"), "護理")
    eq("台照乙 區塊=台籍照服", got.get("台照乙"), "台籍照服")
    eq("外照丙 區塊=外籍照服", got.get("外照丙"), "外籍照服")
    # 區塊對了，照護表的「白班照服」才取得到台籍照服員
    assigns = docgen.namecopy_assignments(conv2, nd)
    eq("照護表2F白班護理=護理甲", assigns["2F"]["1"]["白班護理"], "護理甲")
    eq("照護表5F白班照服=台照乙(只取台籍)", assigns["5F"]["1"]["白班照服"], "台照乙")
    eq("照護表3F白班照服不取外籍", assigns["3F"]["1"]["白班照服"], "")


def test_head_name_same_on_both_paths():
    """迴歸：人頭在兩條路徑上，文件印的姓名必須一致（皆為牌照持有人）。
    路徑一＝T班轉出的 converted；路徑二＝寫出F班再讀回。
    read_fban 若改成取名冊欄姓名，這裡會抓到兩條路徑不一致。"""
    cfg = make_cfg(週起始星期=6)
    cb = CodeBook(CODE_MAP)
    # 路徑一：後台主檔有核章人員(牌照持有人)
    p1 = convert_person(Person(name="陳詡善", block="護理", stamp_name="陳淑萍"),
                        days_from(["Di"] * 31), cb, cfg)
    eq("路徑一 文件用名=陳淑萍", p1["record_name"], "陳淑萍")
    # 路徑二：把同一個人寫進 F 班再讀回
    conv = [_mk_full("陳詡善", "護理", "陳淑萍",
                     {**{d: "D4x" for d in range(1, 32)},
                      **{("fl", d): "2F" for d in range(1, 32)}}, cfg)]
    tmp = _os.path.join(tempfile.gettempdir(), "test_head_paths.xlsx")
    writer.write(conv, 31, cfg, tmp, "115.08")
    conv2, _ = read_fban.load(tmp, cfg)
    p2 = next(x for x in conv2 if x["name"] == "陳詡善")
    eq("路徑二 文件用名=陳淑萍", p2["record_name"], "陳淑萍")
    eq("兩條路徑文件用名一致", p1["record_name"], p2["record_name"])


def run():
    tests = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    for t in tests:
        print(f"\n▶ {t.__name__}")
        try:
            t()
        except Exception:
            global _FAIL; _FAIL += 1; print("  ❌ 例外："); traceback.print_exc()
    print(f"\n{'='*40}\n總計：{_PASS} 通過 / {_FAIL} 失敗")
    return _FAIL


if __name__ == "__main__":
    sys.exit(1 if run() else 0)
