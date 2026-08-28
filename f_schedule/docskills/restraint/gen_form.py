#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""產生「線東大眾護理之家 約束評估記錄單」空白表單（含日期），
並可依班表把三班人員姓名以淺灰色填入簽名欄。一頁=3天x12時段；一張表單(正反面)=6天。"""
import calendar
import copy
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

NAME_FONT = 'KaiTi'
NAME_SIZE = Pt(9)
NAME_GRAY = RGBColor(0xBF, 0xBF, 0xBF)
DATE_FONT = 'KaiTi'
DATE_SIZE = Pt(9)

SLOTS = 12
DAYS_PER_PAGE = 3
DAYS_PER_SHEET = 6


def shift_of(slot_idx):
    if slot_idx < 4:
        return '大夜'
    if slot_idx < 8:
        return '白班'
    return '小夜'


def set_cell_text(cell, text, font, size, color=None, bold=False):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if not text:
        return
    run = p.add_run(text)
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(attr), font)
    if color is not None:
        run.font.color.rgb = color


def zero_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = tcPr.find(qn('w:tcMar'))
    if mar is None:
        mar = tcPr.makeelement(qn('w:tcMar'), {})
        tcPr.append(mar)
    for side in ('top', 'start', 'left', 'bottom', 'end', 'right'):
        el = mar.find(qn('w:' + side))
        if el is None:
            el = mar.makeelement(qn('w:' + side), {})
            mar.append(el)
        el.set(qn('w:w'), '0')
        el.set(qn('w:type'), 'dxa')


def patch_header_line(par, roc_year, month):
    runs = par.runs
    for i, r in enumerate(runs):
        if r.text == '年' and i > 0 and set(runs[i - 1].text) == {'_'}:
            width = len(runs[i - 1].text)
            runs[i - 1].text = str(roc_year).center(width, '_')
        if r.text == '月' and i > 0 and set(runs[i - 1].text) == {'_'}:
            width = len(runs[i - 1].text)
            runs[i - 1].text = ('%d' % month).center(width, '_')


def build(base_path, roc_year, month, out_path, assignments=None):
    """assignments: {day(int): {'大夜':name,'白班':name,'小夜':name}}"""
    year = roc_year + 1911
    ndays = calendar.monthrange(year, month)[1]
    npages = -(-ndays // DAYS_PER_PAGE)
    nsheets = -(-ndays // DAYS_PER_SHEET)
    if npages % 2:
        npages += 1

    doc = Document(base_path)
    body = doc.element.body
    children = list(body.iterchildren())
    group = [copy.deepcopy(ch) for ch in children[0:4]]
    sectPr = children[-1]

    for ch in children:
        if ch is not sectPr:
            body.remove(ch)

    day = 1
    for page in range(npages):
        blk = [copy.deepcopy(ch) for ch in group]
        for ch in blk:
            body.insert(body.index(sectPr), ch)
        tbl_el = blk[1]
        tbl = [t for t in doc.tables if t._tbl is tbl_el][0]
        patch_header_line(_para(doc, blk[3]), roc_year, month)

        for k in range(DAYS_PER_PAGE):
            d = day + k
            for s in range(SLOTS):
                row = tbl.rows[1 + k * SLOTS + s]
                date_cell = row.cells[0]
                sign_cell = row.cells[7]
                zero_cell_margins(date_cell)
                zero_cell_margins(sign_cell)
                if d > ndays:
                    set_cell_text(date_cell, '', DATE_FONT, DATE_SIZE)
                    set_cell_text(sign_cell, '', NAME_FONT, NAME_SIZE)
                    continue
                set_cell_text(date_cell, '%d/%d' % (month, d), DATE_FONT, DATE_SIZE)
                name = ''
                if assignments:
                    name = assignments.get(d, {}).get(shift_of(s), '') or ''
                set_cell_text(sign_cell, name, NAME_FONT, NAME_SIZE, NAME_GRAY)
        day += DAYS_PER_PAGE

    doc.save(out_path)
    return dict(days=ndays, pages=npages, sheets=nsheets)


def _para(doc, p_el):
    from docx.text.paragraph import Paragraph
    return Paragraph(p_el, doc.element.body)


if __name__ == '__main__':
    base, y, m, out = sys.argv[1], int(sys.argv[2]), int(sys.argv[3]), sys.argv[4]
    print(build(base, y, m, out))
