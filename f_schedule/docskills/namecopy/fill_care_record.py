#!/usr/bin/env python3
"""Fill a floor's 住民日常生活照護表 (.docx) with 白班護理/白班照服/夜班照服 names."""
import sys
import json
import argparse
import re
import calendar
import docx
from docx.oxml.ns import qn
from docx.shared import Pt

CATEGORY_BY_ROLE = {
    "aide_day": "白班照服",
    "aide_night": "夜班照服",
    "nurse_day": "白班護理",
}

NAME_FONT = "KaiTi"
NAME_FONT_SIZE_PT = 8
NAME_BOLD = True


def set_run_font(run, name=NAME_FONT, size_pt=NAME_FONT_SIZE_PT, bold=NAME_BOLD):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def lock_table_layout(table):
    table.autofit = False


TITLE_LINE_PATTERN = re.compile(r"\d{2,3}\s*年\s*\d{1,2}\s*月.*姓名")
MAX_RUN_OF_SPACES = 25


def normalize_title_spacing(doc):
    for p in doc.paragraphs:
        if TITLE_LINE_PATTERN.search(p.text):
            for run in p.runs:
                run.text = re.sub(r" {%d,}" % (MAX_RUN_OF_SPACES + 1),
                                  " " * MAX_RUN_OF_SPACES, run.text)


def set_document_font(doc, name="KaiTi"):
    def set_paragraph_font(p):
        for run in p.runs:
            run.font.name = name
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rPr.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), name)
            rFonts.set(qn("w:ascii"), name)
            rFonts.set(qn("w:hAnsi"), name)
    for p in doc.paragraphs:
        set_paragraph_font(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_font(p)
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            set_paragraph_font(p)
        for p in sec.footer.paragraphs:
            set_paragraph_font(p)


def detect_days_in_month(doc):
    pattern = re.compile(r"(\d{2,3})\s*年\s*(\d{1,2})\s*月")
    for p in doc.paragraphs:
        m = pattern.search(p.text)
        if m:
            western_year = int(m.group(1)) + 1911
            try:
                return calendar.monthrange(western_year, int(m.group(2)))[1]
            except ValueError:
                return None
    return None


def clear_cell(cell):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ""
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)


def blank_out_of_range_days(table, day_map, days_in_month):
    if days_in_month is None:
        return
    out_of_range_cols = {col for day, col in day_map.items() if day > days_in_month}
    if not out_of_range_cols:
        return
    for row in table.rows:
        for col_idx in out_of_range_cols:
            if col_idx < len(row.cells):
                clear_cell(row.cells[col_idx])


def _row_height_val(row):
    trPr = row._tr.find(qn("w:trPr"))
    if trPr is None:
        return 0
    h = trPr.find(qn("w:trHeight"))
    if h is None:
        return 0
    v = h.get(qn("w:val"))
    return int(v) if v and v.isdigit() else 0


def set_row_height(row, val):
    """把列高設為至少 val（twips），避免責任護士列比照服員列小。"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = tr.makeelement(qn("w:trPr"), {})
        tr.insert(0, trPr)
    h = trPr.find(qn("w:trHeight"))
    if h is None:
        h = trPr.makeelement(qn("w:trHeight"), {})
        trPr.append(h)
    h.set(qn("w:val"), str(val))
    h.set(qn("w:hRule"), "atLeast")


def equalize_name_rows(table, roles):
    """責任護士列高度＝各姓名列的最大值，讓它不比照服員列小。"""
    rows = [table.rows[ri] for ri in roles.values()]
    if not rows:
        return
    target = max([_row_height_val(r) for r in rows] + [260])
    nd = roles.get("nurse_day")
    if nd is not None:
        set_row_height(table.rows[nd], target)


def find_special_rows(table):
    roles = {}
    for ri, row in enumerate(table.rows):
        c0 = row.cells[0].text.strip()
        if "照服員" in c0:
            c2 = row.cells[2].text.strip() if len(row.cells) > 2 else ""
            if c2 == "白":
                roles["aide_day"] = ri
            elif c2 in ("晚", "夜"):
                roles["aide_night"] = ri
        elif "責任護士" in c0:
            roles["nurse_day"] = ri
    return roles


def build_day_column_map(table):
    header = table.rows[0]
    day_map = {}
    for ci, cell in enumerate(header.cells):
        text = cell.text.strip()
        if text.isdigit():
            day = int(text)
            if day not in day_map:
                day_map[day] = ci
    return day_map


def set_cell_name(cell, name):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ""
    if p.runs:
        p.runs[0].text = name
        set_run_font(p.runs[0])
    else:
        set_run_font(p.add_run(name))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)


def fill(template_path, assignments_path, floor, output_path):
    with open(assignments_path, "r", encoding="utf-8") as f:
        assignments = json.load(f)
    if floor not in assignments:
        raise SystemExit(f"Floor '{floor}' not found (available: {list(assignments.keys())})")
    floor_data = assignments[floor]

    d = docx.Document(template_path)
    if not d.tables:
        raise SystemExit("No tables found in the template docx.")

    set_document_font(d, "KaiTi")
    normalize_title_spacing(d)

    days_in_month = detect_days_in_month(d)
    if days_in_month is None:
        print("WARNING: could not detect '<年>年 <月>月'", file=sys.stderr)
    else:
        print(f"Detected {days_in_month} day(s) in this month.")

    filled_days = set()
    for table in d.tables:
        roles = find_special_rows(table)
        if not roles:
            continue
        day_map = build_day_column_map(table)
        if not day_map:
            continue
        blank_out_of_range_days(table, day_map, days_in_month)
        lock_table_layout(table)
        equalize_name_rows(table, roles)
        # 每個姓名列先整列清空（日期欄起），再只填當月有效日；
        # 如此超出當月天數的欄位（如 9 月的第 31 欄、或範本右側多出的空欄）一律留白。
        first_day_col = min(day_map.values())
        for role, row_idx in roles.items():
            category = CATEGORY_BY_ROLE[role]
            row = table.rows[row_idx]
            for ci in range(first_day_col, len(row.cells)):
                clear_cell(row.cells[ci])
            for day, col_idx in day_map.items():
                if days_in_month is not None and day > days_in_month:
                    continue
                name = floor_data.get(str(day), {}).get(category)
                if not name or col_idx >= len(row.cells):
                    continue
                set_cell_name(row.cells[col_idx], name)
                filled_days.add(day)

    d.save(output_path)
    print(f"Saved {output_path}. Filled data for {len(filled_days)} day(s).")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("template_docx")
    ap.add_argument("assignments_json")
    ap.add_argument("floor", choices=["2F", "3F", "5F"])
    ap.add_argument("output_docx")
    args = ap.parse_args()
    fill(args.template_docx, args.assignments_json, args.floor, args.output_docx)


if __name__ == "__main__":
    main()
